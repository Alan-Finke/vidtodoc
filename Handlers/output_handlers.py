import shutil
from typing import Optional
from docx import Document
from Utilities.docx_conversions import convert_docx_to_pdf_libreoffice
from Utilities.html_writer import HTMLWriter
from Utilities.markdown_writer import MarkdownWriter
import os

class BaseOutputHandler:
    """Abstract base class for output handlers."""
    def add_title(self, title: str): pass
    def add_summary(self, summary: str): pass
    def add_steps_heading(self): pass
    def add_step(self, text: str, frame_file: str, counter: int): pass
    def save(self, output_path: str): pass
    def postprocess(self, output_path: str): pass
    def cleanup(self, output_path: str, clean_frames: bool = False):
        base_name = os.path.basename(output_path)
        root_frames_dir = f"{base_name.split('.')[0]}_{base_name.split('.')[1]}"
        if clean_frames and os.path.exists(root_frames_dir):
            # Clean up the frame folder
            shutil.rmtree(root_frames_dir)
        # Traverse through subfolders and delete all __pycache__ directories
        for root, dirs, files in os.walk(os.getcwd()):
            if "__pycache__" in dirs:
                shutil.rmtree(os.path.join(root, "__pycache__"))

class DocxOutputHandler(BaseOutputHandler):
    """Handles DOC and DOCX output."""
    def __init__(self):
        self.document = Document()
    def add_title(self, title: str):
        self.document.add_heading(title, level=1)
    def add_summary(self, summary: str):
        self.document.add_heading("Summary", level=2)
        self.document.add_paragraph(summary)
    def add_steps_heading(self):
        self.document.add_heading("Steps", level=2)
    def add_step(self, text: str, frame_file: str, counter: int, image_width: float):
        self.document.add_paragraph(text)
        self.document.add_picture(frame_file, width=image_width)  # Requires 'docx.shared.Inches'
    def save(self, output_path: str):
        self.document.save(output_path)
    def postprocess(self, output_path: str):
        self.cleanup(output_path, True)
        
class PdfOutputHandler(DocxOutputHandler):
    """Handles PDF output via DOCX conversion."""
    def save(self, output_path: str):
        temp_docx = output_path.replace('.pdf', '.docx')
        self.document.save(temp_docx)
        convert_docx_to_pdf_libreoffice(temp_docx, output_path)
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
    def postprocess(self, output_path):
        self.cleanup(output_path, True)

class MarkdownOutputHandler(BaseOutputHandler):
    """Handles Markdown output."""
    def __init__(self):
        self.md_document = MarkdownWriter()
    def add_title(self, title: str):
        self.md_document.add_heading(title, level=1)
    def add_summary(self, summary: str):
        self.md_document.add_heading("Summary", level=2)
        self.md_document.add_paragraph(summary)
    def add_steps_heading(self):
        self.md_document.add_heading("Steps", level=2)
    def add_step(self, text: str, frame_file: str, counter: int, image_width: float):
        self.md_document.add_paragraph(text)
        self.md_document.add_image(frame_file, image_width, f"Image {counter}")
    def save(self, output_path: str):
        self.md_document.save(output_path)
    def postprocess(self, output_path: str):
        self.cleanup(output_path)

class HtmlOutputHandler(BaseOutputHandler):
    """Handles HTML output using a template."""
    def __init__(self, templates_path: str):
        self.html_document = HTMLWriter(templates_path)
    def add_title(self, title: str):
        self.html_document.add_title(title)
    def add_summary(self, summary: str):
        self.html_document.add_summary(summary)
    def add_steps_heading(self):
        pass  # handled in template
    def add_step(self, text: str, frame_file: str, counter: int, image_width: float):
        self.html_document.add_step(counter, text, frame_file, image_width, alt_text=f"Image {counter}")
    def save(self, output_path: str):
        self.html_document.save(output_path)
    def postprocess(self, output_path: str):
        self.cleanup(output_path)

# Registry for output handlers
OUTPUT_HANDLER_REGISTRY = {
    "doc": DocxOutputHandler,
    "docx": DocxOutputHandler,
    "pdf": PdfOutputHandler,
    "md": MarkdownOutputHandler,
    "html": HtmlOutputHandler
}

def get_output_handler(output_format: str, template_path: str) -> BaseOutputHandler:
    """
    Factory function to get the appropriate output handler.
    """
    if output_format not in OUTPUT_HANDLER_REGISTRY:
        raise ValueError(f"Unsupported output format: {output_format}")
    return OUTPUT_HANDLER_REGISTRY[output_format](template_path) if output_format == "html" else OUTPUT_HANDLER_REGISTRY[output_format]()