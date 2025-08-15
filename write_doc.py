import os
from docx import Document
from docx.shared import Inches
import extract_frame

def write_doc_file(title, summary, vidfile, result):
    # Writes the content of a document to a DOCX file.
    print("Writing DOCX file...")

    # Create a new Document
    document = Document()
    
    # Add a title
    document.add_heading(title, level=1)

    # Add summary section
    document.add_heading("Summary", level=2)
    document.add_paragraph(summary)

    ### Add the Steps section
    document.add_heading("Steps", level=2)
    for segment in result["segments"]:
        print(f"[{segment['start']:.2f} - {segment['end']:.2f}] {segment['text']}")
        # Extract a frame at the start of each segment
        extract_frame.extract_frame(vidfile, dir, segment['start']);

        #Add the text to the document
        document.add_paragraph(segment['text']) ;

        # Add the frame to the document
        document.add_picture("frame.jpg", width=Inches(5.00));  # Requires 'docx.shared.Inches'


    ### Close and save the document
    document.save("MyVideo.docx");

